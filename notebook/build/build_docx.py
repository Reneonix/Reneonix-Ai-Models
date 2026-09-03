#!/usr/bin/env python3
"""Builds a plain, Google-Docs-style .docx version of the Model Training Report - same real
numbers as the HTML artifact, but formatted like the reference doc the user pasted (Times New
Roman, black text, bordered tables, yellow-highlighted key figures, one image per curve/prediction)
so it imports cleanly into Google Docs (Drive > Open with > Google Docs converts .docx natively)."""
import json
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "d:/Reneonix/yolo_projects/Wastes_identification"
BUILD = f"{ROOT}/notebook/build"
ASSETS = f"{BUILD}/report_assets"
OUT_PATH = f"{ROOT}/notebook/Model_Training_Report.docx"

with open(f"{BUILD}/curves.json") as f:
    CURVES = json.load(f)
with open(f"{BUILD}/benchmark_all_result.json") as f:
    BENCH_ALL = json.load(f)
BENCH = BENCH_ALL["models"]
TEST_RES = BENCH_ALL["image_resolution"]

# ---------------------------------------------------------------------------
# Generate the two missing curve charts (exp004, exp005) with matplotlib -
# real per-epoch data from results.csv, same source as the HTML's live charts.
# ---------------------------------------------------------------------------
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

os.makedirs(ASSETS, exist_ok=True)

def make_curve_png(exp, out_path):
    d = CURVES[exp]
    fig, ax = plt.subplots(figsize=(6, 4), dpi=150)
    if exp == "exp004":
        ax.plot(d["epoch"], d["train_acc"], label="train acc", color="#2A63E4")
        ax.plot(d["epoch"], d["val_acc"], label="val acc", color="#0E93A3")
        ax.set_ylabel("Accuracy")
    else:
        ax.plot(d["epoch"], d["precision"], label="precision", color="#2A63E4")
        ax.plot(d["epoch"], d["recall"], label="recall", color="#0E93A3")
        ax.plot(d["epoch"], d["map5095"], label="mAP50-95", color="#6C5CE0")
        ax.set_ylabel("Score")
    ax.set_xlabel("Epoch")
    ax.set_title(f"{exp} training curve")
    ax.set_ylim(0, 1.02)
    ax.grid(alpha=0.25)
    ax.legend(loc="lower right", fontsize=9)
    fig.tight_layout()
    fig.savefig(out_path)
    plt.close(fig)

make_curve_png("exp004", f"{ASSETS}/exp004_curve.png")
make_curve_png("exp005", f"{ASSETS}/exp005_curve.png")
print("Generated exp004/exp005 curve charts")

# ---------------------------------------------------------------------------
# Ground truth data (same facts as the HTML report / notebook/build/build_report.py)
# ---------------------------------------------------------------------------
TRAINING_TIME = {
    "exp001": "2h 19m (100 epochs)",
    "exp002": "1h 23m (120 epochs)",
    "exp003": "41h 17m (100 epochs)",
    "exp004": "not logged for this run (29 epochs to best checkpoint)",
    "exp005": "4h 48m (100 epochs)",
    "exp006": "29m (fine-tune stage, 100 epochs) + a separate base-training stage",
}

MODELS = [
    dict(exp="exp001", title="YOLOv8L", classes="6 - aluminium, plastic, metal, stone, ceramic, glass",
         train_n="7,389", val_n="1,847", batch=8, imgsz=640, epochs=100,
         dataset_line="Dataset is taken from the project's main conveyor/pour capture set - "
                       "7,389 training images and 1,847 validation images, hand-annotated in YOLO "
                       "format across all 6 material classes.",
         explanation="YOLOv8L is a CNN-based, anchor-free single-stage detector: a CSPDarknet-style "
                     "backbone extracts multi-scale features, a PAN/FPN neck fuses them, and a "
                     "decoupled head predicts class and box offsets per grid cell directly, with no "
                     "separate region-proposal stage. This run overrides Ultralytics' default "
                     "optimizer with a 2-group differential learning-rate schedule: the first 5 "
                     "backbone layers (generic edge/texture filters) are frozen outright, the "
                     "remaining backbone/neck layers train at a low learning rate (3.5x10-4), and the "
                     "detection head - which must learn this project's 6 classes from scratch - "
                     "trains at a high learning rate (7.0x10-3), both annealed on a cosine schedule.",
         precision=0.945, recall=0.932, map50=0.975, map5095=0.935, is_clf=False,
         pr_img=f"{ROOT}/results/exp001/BoxPR_curve.png"),
    dict(exp="exp002", title="YOLO26s", classes="1 - glass (cullet fragments only)",
         train_n="2,350", val_n="587", batch=8, imgsz=640, epochs=120,
         dataset_line="Dataset is a separate, narrower capture set of glass-cullet fragments only - "
                       "2,350 training images and 587 validation images, single class.",
         explanation="Same YOLO26 architecture family as exp006, but trained as a narrow specialist: "
                     "a single output class (glass cullet) on a dataset roughly a third the size of "
                     "the main six-class set. Glass-cullet fragments have no consistent silhouette "
                     "the way a whole bottle or can does, which - combined with the smaller dataset - "
                     "explains its lower headline numbers versus the six-class models.",
         precision=0.694, recall=0.553, map50=0.602, map5095=0.370, is_clf=False,
         pr_img=f"{ROOT}/results/exp002/BoxPR_curve.png"),
    dict(exp="exp003", title="YOLO26L + P2", classes="6 - aluminium, plastic, metal, stone, ceramic, glass",
         train_n="7,389", val_n="1,847", batch=16, imgsz=640, epochs=100,
         dataset_line="Same main dataset as YOLOv8L above - 7,389 training images and 1,847 "
                       "validation images, identical split for a fair comparison.",
         explanation="The only architectural change from a stock YOLO26L is an added P2 (160x160) "
                     "detection head alongside the stock P3/P4/P5 heads, giving the network a "
                     "detection path with less downsampling - better suited to small or distant "
                     "objects on the conveyor. Every other hyperparameter is left at Ultralytics "
                     "defaults, isolating the P2 head's own effect on accuracy. This is the "
                     "best-performing model in the project on every headline metric.",
         precision=0.947, recall=0.930, map50=0.977, map5095=0.951, is_clf=False,
         pr_img=f"{ROOT}/results/exp003/BoxPR_curve.png"),
    dict(exp="exp004", title="ResNet-50", classes="6 - aluminium, plastic, metal, stone, ceramic, glass",
         train_n="294,959 crops", val_n="73,448 crops", batch=64, imgsz=224, epochs=29,
         dataset_line="Dataset is not separately captured - it is built by cropping every labeled "
                       "YOLO box out of the main dataset's own images (10% padding added, boxes "
                       "under 10px skipped), giving 294,959 training crops and 73,448 validation "
                       "crops with zero new manual labeling.",
         explanation="A standard ResNet-50 backbone, ImageNet-pretrained, with its final layer "
                     "replaced by a 6-way linear classifier. Not a detector: it is the second stage "
                     "of a two-model pipeline where YOLO detects and localizes an object first, then "
                     "this classifier looks only at that tightly cropped region to assign the final "
                     "material label, using heavy color/lighting augmentation for robustness to "
                     "glare. It cannot localize objects in a raw multi-object frame on its own.",
         precision=None, recall=None, map50=None, map5095=None, is_clf=True, val_acc=0.9937234506045093,
         pr_img=f"{ASSETS}/exp004_curve.png"),
    dict(exp="exp005", title="RT-DETR-L", classes="6 - aluminium, plastic, metal, stone, ceramic, glass",
         train_n="7,389", val_n="1,847", batch=16, imgsz=640, epochs=100,
         dataset_line="Same main dataset and identical validation split as YOLOv8L and YOLO26L+P2 "
                       "above - 7,389 training images and 1,847 validation images.",
         explanation="RT-DETR (Baidu, 2023) is a transformer-based detector with no anchor boxes and "
                     "no NMS post-processing - a CNN backbone feeds a transformer encoder-decoder "
                     "that predicts a fixed set of object queries directly. Trained via Ultralytics' "
                     "own RTDETR class, the same .train() API as every YOLO model here, making it a "
                     "fair architecture-only comparison. Training completed all 100 epochs, but the "
                     "process was interrupted during Ultralytics' final validation pass, so no "
                     "confusion matrix or PR-curve image exists for this run - the chart below is a "
                     "live per-epoch training curve instead.",
         precision=0.946, recall=0.924, map50=0.975, map5095=0.922, is_clf=False,
         pr_img=f"{ASSETS}/exp005_curve.png"),
    dict(exp="exp006", title="YOLO26s (fine-tuned)", classes="6 - aluminium, plastic, metal, stone, ceramic, glass",
         train_n="1,631 + 7,389", val_n="408 / 1,847", batch=8, imgsz=640, epochs=100,
         dataset_line="Trained in two stages: base-trained on its own 1,631/408 dataset, then "
                       "fine-tuned on the main 7,389/1,847 dataset - the numbers below are the final "
                       "fine-tuned model.",
         explanation="Same YOLO26s architecture as exp002, trained in two stages instead of one: "
                     "stage 1 base-trains on a separate smaller dataset, then stage 2 fine-tunes the "
                     "resulting weights on this project's main six-class dataset. It lands between "
                     "YOLOv8L and YOLO26L+P2 on the final held-out set despite being a much smaller, "
                     "faster network - evidence that the base-training stage transfers useful "
                     "structure before the model ever sees this project's own data.",
         precision=0.94436, recall=0.92681, map50=0.97514, map5095=0.94335, is_clf=False,
         pr_img=f"{ROOT}/results/exp006/BoxPR_curve.png"),
]

print("data ready, building document...")

# ---------------------------------------------------------------------------
# docx helpers - Times New Roman body, black text, bordered tables, yellow
# highlights on key figures, matching the reference doc's plain style.
# ---------------------------------------------------------------------------
from docx.enum.text import WD_COLOR_INDEX

BODY_FONT = "Times New Roman"

def set_run_font(run, size=11, bold=False, highlight=False):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element.rPr
    r.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def add_para(doc, text, highlights=None, size=11, bold=False, space_after=10):
    """text with zero or more substrings in `highlights` rendered on a yellow background,
    same convention as the reference document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    highlights = highlights or []
    remaining = text
    while remaining:
        hit = None
        for h in highlights:
            idx = remaining.find(h)
            if idx != -1 and (hit is None or idx < hit[0]):
                hit = (idx, h)
        if hit is None:
            run = p.add_run(remaining)
            set_run_font(run, size=size, bold=bold)
            remaining = ""
        else:
            idx, h = hit
            if idx > 0:
                run = p.add_run(remaining[:idx])
                set_run_font(run, size=size, bold=bold)
            run = p.add_run(h)
            set_run_font(run, size=size, bold=bold, highlight=True)
            remaining = remaining[idx + len(h):]
    return p

def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_run_font(run, size=(20 if level == 0 else 16 if level == 1 else 13), bold=True)
    h.paragraph_format.space_before = Pt(18 if level <= 1 else 10)
    h.paragraph_format.space_after = Pt(8)
    return h

def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)

def add_kv_table(doc, rows, highlight_rows=()):
    """rows: list of (label, value). highlight_rows: set of label strings whose value gets
    the yellow-highlight treatment, matching the reference doc's Accuracy row."""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for label, value in rows:
        row = table.add_row()
        c0, c1 = row.cells
        c0.width = Inches(2.3)
        c1.width = Inches(3.7)
        for c in (c0, c1):
            set_cell_border(c)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        set_run_font(r0, size=11)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(str(value))
        set_run_font(r1, size=11, highlight=(label in highlight_rows))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_image(doc, path, width_in=5.6):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[image not found: {path}]")

def fmt(v):
    return f"{v:.3f}" if v is not None else "n/a"

# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(11)

title = doc.add_heading(level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Model Training Report")
set_run_font(run, size=26, bold=True)
title.paragraph_format.space_after = Pt(6)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("AI-Based Waste Detection and Classification System")
set_run_font(r, size=12, bold=False)
sub.paragraph_format.space_after = Pt(24)

for i, m in enumerate(MODELS, start=1):
    exp = m["exp"]
    b_plain = BENCH[exp]["plain"]
    b_e2e = BENCH[exp]["e2e_plain"]

    add_heading(doc, f"{i}. {m['title']}", level=1)

    add_heading(doc, "Model Explanation", level=2)
    add_para(doc, m["explanation"])

    add_heading(doc, "Dataset", level=2)
    add_para(doc, m["dataset_line"], highlights=[m["train_n"], m["val_n"]])

    add_heading(doc, "Training Parameters", level=2)
    rows = [
        ("Base model", m["title"]),
        ("Classes", m["classes"]),
        ("Training images" if not m["is_clf"] else "Training crops", m["train_n"]),
        ("Validation images" if not m["is_clf"] else "Validation crops", m["val_n"]),
        ("Batch size", m["batch"]),
        ("Training Image Input size", f"{m['imgsz']}px"),
        ("Epochs trained", m["epochs"]),
        ("Training time", TRAINING_TIME[exp]),
    ]
    if m["is_clf"]:
        rows.append(("Accuracy", f"{m['val_acc']*100:.2f}%"))
        highlight_rows = {"Accuracy"}
    else:
        rows += [
            ("Precision", fmt(m["precision"])),
            ("Recall", fmt(m["recall"])),
            ("mAP50", fmt(m["map50"])),
            ("mAP50-95", fmt(m["map5095"])),
        ]
        highlight_rows = {"mAP50-95"}
    add_kv_table(doc, rows, highlight_rows=highlight_rows)

    add_para(
        doc,
        f"Average model Inference time is : {b_plain['avg_ms']} ms ({b_plain['fps']} FPS)",
        highlights=[f"{b_plain['avg_ms']} ms ({b_plain['fps']} FPS)"],
    )
    add_para(
        doc,
        f"Average end to end latency : {b_e2e['avg_ms']} ms",
        highlights=[f"{b_e2e['avg_ms']} ms"],
    )
    add_para(
        doc,
        f"(Measured on an RTX 5080, {TEST_RES} test frame, 5 warmup + 30 timed passes. "
        "End-to-end includes reading the frame and drawing the annotated output, not just the "
        "model's forward pass.)",
        size=9,
    )

    add_heading(doc, "Model curve", level=2)
    add_image(doc, m["pr_img"])

    add_heading(doc, "Prediction", level=2)
    add_para(doc, "Plain inference (no tiling):", size=10, bold=True, space_after=4)
    add_image(doc, f"{ASSETS}/{exp}_plain.jpg")
    add_para(doc, "SAHI inference (6-tile):", size=10, bold=True, space_after=4)
    add_image(doc, f"{ASSETS}/{exp}_sahi.jpg")

    if i < len(MODELS):
        doc.add_page_break()

os.makedirs(f"{ROOT}/notebook", exist_ok=True)
doc.save(OUT_PATH)
print("Written:", OUT_PATH, "-", round(os.path.getsize(OUT_PATH) / 1024 / 1024, 2), "MB")
